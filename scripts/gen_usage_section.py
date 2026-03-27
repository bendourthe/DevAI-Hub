"""Generate a styled Word document with the comprehensive Usage Limit Management section."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

HEADER_FILL = "111D35"
ALT_ROW_FILL = "F0F4F8"
WHITE = "FFFFFF"

SOURCE = r"C:\Users\BEDOURTHE\Downloads\AI Training Program\05 Agentic Coding Handbook.docx"
OUTPUT = r"C:\Users\BEDOURTHE\Downloads\AI Training Program\Usage Limit Management Section.docx"


def set_cell_shading(cell, color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    tcPr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False, font_color: str | None = None, font_size=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_color:
        run.font.color.rgb = RGBColor.from_string(font_color)
    if font_size:
        run.font.size = font_size


def add_styled_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Normal Table"

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, h, bold=True, font_color=WHITE, font_size=Pt(10))

    for r_idx, row_data in enumerate(rows):
        fill = WHITE if r_idx % 2 == 0 else ALT_ROW_FILL
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_shading(cell, fill)
            set_cell_text(cell, val, font_size=Pt(10))

    return table


def add_bold_normal(doc, bold_part: str, rest: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(bold_part)
    r.bold = True
    p.add_run(rest)


def add_numbered(doc, bold_part: str, rest: str = "") -> None:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(bold_part)
    r.bold = True
    if rest:
        p.add_run(rest)


def add_bullet(doc, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def build_document() -> None:
    doc = Document(SOURCE)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    for t in doc.tables:
        t._element.getparent().remove(t._element)

    # =========================================================================
    # USAGE LIMIT MANAGEMENT
    # =========================================================================
    doc.add_heading("Usage Limit Management", level=1)

    doc.add_paragraph(
        "Both Claude Code and Codex operate on tiered usage limits that cap how much you can use within a given time window. "
        "Understanding these limits, monitoring your consumption, and applying optimization strategies prevents unexpected interruptions and helps you get the most out of your AI coding budget."
    )

    # --- Understanding Usage Limits ---
    doc.add_heading("Understanding Usage Limits", level=2)

    doc.add_paragraph(
        "Each platform enforces multiple independent limits. Hitting one limit does not affect the others. "
        "Limits are based on rolling time windows (not calendar days), so they reset continuously rather than at midnight."
    )

    doc.add_heading("Claude Code Limits", level=3)

    doc.add_paragraph(
        "Claude Code enforces three independent limits. Each operates on its own rolling window and resets independently."
    )

    add_styled_table(doc,
        ["Limit", "Scope", "Reset Cycle", "Primary Drain"],
        [
            [
                "Current Session",
                "Token budget per conversation. Covers input tokens, output tokens, and tool call overhead.",
                "Rolling 5-hour window.",
                "Large agentic runs with many tool calls, extensive file reads, and long thinking chains.",
            ],
            [
                "Weekly (All Models)",
                "Rolling 7-day cap across all models (Opus, Sonnet, Haiku).",
                "Every 7 days from first trigger.",
                "Sustained heavy use across multiple sessions and days.",
            ],
            [
                "Weekly (Sonnet Only)",
                "Separate cap specifically on Sonnet model usage. Independent from the all-models cap.",
                "Independent rolling 7-day window.",
                "Developers who use Sonnet as their primary model for extended coding sessions.",
            ],
        ]
    )

    doc.add_paragraph(
        "The session limit is the one you will hit most frequently during active development. When it triggers, you can either wait for the 5-hour window to roll forward or start a new conversation (which does not reset the limit but does start fresh context)."
    )

    doc.add_heading("Codex Limits", level=3)

    doc.add_paragraph(
        "Codex uses a dual-limit system where local messages and cloud tasks share the same quota bucket within each window."
    )

    add_styled_table(doc,
        ["Limit", "Scope", "Reset Cycle", "Primary Drain"],
        [
            [
                "5-Hour Window",
                "Combined quota for local messages and cloud tasks. Heavy local message use reduces available cloud task capacity.",
                "Rolling 5-hour window.",
                "Frequent prompt submissions and cloud task dispatches within a single work block.",
            ],
            [
                "Weekly Quota",
                "Cumulative cap across all 5-hour windows during a 7-day cycle.",
                "Rolling 7-day window.",
                "Sustained daily use across the full work week.",
            ],
        ]
    )

    doc.add_paragraph(
        "Approximate quotas vary by plan tier. Plus plans allow roughly 30 to 150 messages per 5-hour window and approximately 3,000 requests per week. "
        "Pro and Enterprise tiers have significantly higher quotas with faster resets and team-based credit pools. "
        "A single complex prompt can consume 5 to 7% of the weekly quota, so prompt efficiency matters."
    )

    # --- Monitoring Your Usage ---
    doc.add_heading("Monitoring Your Usage", level=2)

    doc.add_paragraph(
        "DevAI-Hub provides three complementary monitoring tools for Claude Code, each suited to a different workflow. "
        "Use them together for complete visibility."
    )

    doc.add_heading("Claude Usage Monitor: VS Code Extension", level=3)

    doc.add_paragraph(
        "The Claude Usage Monitor is a VS Code extension that provides real-time usage visibility directly in your editor. "
        "It is the most comprehensive monitoring tool available and the recommended way to track your Claude Code consumption."
    )

    add_bold_normal(doc, "Installation: ", "The extension is included in the DevAI-Hub repository under extensions/claude-usage-monitor/.")

    add_numbered(doc, "Open a terminal in the extension directory: ", "cd extensions/claude-usage-monitor")
    add_numbered(doc, "Install dependencies: ", "npm install")
    add_numbered(doc, "Compile the extension: ", "npm run compile")
    add_numbered(doc, "Package and install: ", "npm run package (creates a .vsix file), then install it in VS Code via Extensions > Install from VSIX.")

    doc.add_paragraph(
        "Alternatively, press F5 in VS Code while the extension folder is open to launch it in development mode."
    )

    add_bold_normal(doc, "Status bar: ", "Once installed, the extension adds a Claude Usage indicator to the VS Code status bar. It displays your current session and weekly usage percentages at a glance. The background is color-coded: green (0 to 50%), yellow (51 to 75%), and red (76 to 100%).")

    add_bold_normal(doc, "Hover tooltip: ", "Hover over the status bar item to see a detailed tooltip with SVG progress bars for each metric (session, weekly all-models, weekly Sonnet-only, Opus usage) and countdown timers showing when each limit resets.")

    add_bold_normal(doc, "Dashboard: ", "Click the status bar item to open a full usage dashboard panel. The dashboard shows all metrics, provides model-switching recommendations based on your current consumption, and offers optimization tips tailored to your usage pattern.")

    add_bold_normal(doc, "Auto-switcher: ", "The extension can automatically switch your Claude Code model when usage exceeds configurable thresholds. For example, it can switch from Opus to Sonnet when weekly usage hits 75%, and from Sonnet to Haiku when it hits 95%. This prevents you from hitting hard limits unexpectedly.")

    doc.add_paragraph("Key settings (accessible via VS Code Settings, search \"Claude Usage\"):")

    add_styled_table(doc,
        ["Setting", "Default", "Description"],
        [
            ["claudeUsage.autoFetch", "true", "Automatically fetch usage data on startup and at intervals."],
            ["claudeUsage.refreshInterval", "10 minutes", "How often to auto-refresh usage data (range: 5 to 120 minutes)."],
            ["claudeUsage.showInStatusBar", "true", "Show or hide the status bar indicator."],
            ["claudeUsage.autoSwitch.enabled", "true", "Enable automatic model switching when thresholds are exceeded."],
            ["claudeUsage.autoSwitch.modelSonnetThreshold", "75%", "Weekly usage % that triggers an Opus to Sonnet switch."],
            ["claudeUsage.autoSwitch.modelHaikuThreshold", "95%", "Weekly usage % that triggers a Sonnet to Haiku switch."],
        ]
    )

    doc.add_paragraph(
        "The extension reads your OAuth token from ~/.claude/.credentials.json (written automatically by Claude Code during authentication). "
        "No additional API keys or configuration are required."
    )

    doc.add_heading("CLI Usage Display (Automatic)", level=3)

    doc.add_paragraph(
        "A Stop hook that fires after every Claude Code response, displaying a compact one-line usage summary directly in the terminal. "
        "It is silent when all metrics are below 50% and only appears when usage becomes elevated. "
        "The display is color-coded using the same green/yellow/orange/red scheme."
    )

    doc.add_paragraph(
        "Example output: Usage: Session 45% | Weekly 82% | Sonnet 12%  (Weekly resets in 2h 15m)"
    )

    doc.add_paragraph(
        "This hook is installed automatically by the DevAI-Hub installer. It requires curl and jq. "
        "Usage data is cached for 5 minutes to avoid excessive API calls."
    )

    doc.add_heading("/check-usage Command (On-Demand)", level=3)

    doc.add_paragraph(
        "For a detailed, on-demand usage report with model-switching recommendations, run /check-usage in any Claude Code session. "
        "The command auto-fetches live data from the Anthropic API and produces a Markdown table with all three metrics, their current percentages, reset timers, and urgency classifications. "
        "If the API is unavailable (for example, if credentials have expired), it falls back to prompting you for manual input from the claude.ai/settings/usage page."
    )

    doc.add_paragraph(
        "The command also generates specific optimization tips based on your current model and usage level, "
        "such as recommending a model switch, suggesting context compression, or advising you to wait for a reset."
    )

    # --- Optimization Strategies ---
    doc.add_heading("Optimization Strategies", level=2)

    doc.add_paragraph(
        "The strategies below apply to both Claude Code and Codex unless noted otherwise. "
        "Applying even a few of these consistently can extend your effective usage by 30 to 50%."
    )

    doc.add_heading("Choose the Right Model", level=3)

    doc.add_paragraph(
        "Not every task requires the most powerful (and most expensive) model. Match the model to the complexity of the task:"
    )

    add_styled_table(doc,
        ["Model", "Best For", "Relative Cost"],
        [
            [
                "Opus (Claude) / o3 (Codex)",
                "Complex reasoning, system architecture, multi-file refactors, nuanced code review.",
                "Highest",
            ],
            [
                "Sonnet (Claude) / o4-mini (Codex)",
                "Feature implementation, bug fixes, test generation, standard code review.",
                "~5x cheaper than Opus",
            ],
            [
                "Haiku (Claude)",
                "Formatting, docstrings, simple lookups, quick Q&A, file renaming.",
                "Near-instant, lowest cost",
            ],
        ]
    )

    doc.add_paragraph(
        "Default to the mid-tier model (Sonnet or o4-mini) for everyday development. "
        "Switch up to Opus/o3 only for tasks that genuinely require deep reasoning. "
        "Switch down to Haiku for trivial tasks. "
        "The Claude Usage Monitor's auto-switcher can handle this transition automatically based on your usage thresholds."
    )

    doc.add_heading("Minimize Context Consumption", level=3)

    doc.add_paragraph("Every token of context the agent processes counts against your limits. Reduce unnecessary context with these strategies:")

    add_bullet(doc, "Point the agent at specific files or directories, not the entire repository. Every 1,000 tokens of additional context increases consumption by approximately 15%.", "Scope your requests: ")
    add_bullet(doc, "Use instruction files (CLAUDE.md, AGENTS.md, GEMINI.md) to pre-define project rules, conventions, and architecture once. The agent reads these automatically instead of you repeating them in every prompt.", "Pre-define project rules: ")
    add_bullet(doc, "Start new conversations when switching to unrelated tasks. Old context from a previous task wastes tokens and can confuse the agent.", "Fresh sessions for new tasks: ")
    add_bullet(doc, "In Claude Code, use /compact to compress the context window mid-session. This summarizes older messages, freeing token budget for new work.", "Compress mid-session (Claude Code): ")
    add_bullet(doc, "Each active MCP server adds context overhead. Disable servers you are not actively using.", "Limit active MCP servers: ")
    add_bullet(doc, "Set the environment variable MAX_THINKING_TOKENS=10000 to cap extended thinking budget. Set CLAUDE_AUTOCOMPACT_PCT_OVERRIDE=50 to trigger automatic context compaction at 50% instead of the default 85%.", "Environment tuning (Claude Code): ")

    doc.add_heading("Use Commands Instead of Prose", level=3)

    doc.add_paragraph(
        "A single /command (or its equivalent workflow in Gemini or custom prompt in Codex) replaces paragraphs of repeated instructions. "
        "Commands are pre-optimized for token efficiency because they encode complex multi-step workflows in a compact format that the agent already understands."
    )

    doc.add_paragraph(
        "For example, /generate-commit-message costs 2 to 5K tokens. Manually describing the same task "
        "(\"analyze my staged changes, categorize them, write a conventional commit message with scope and body\") "
        "costs 3 to 5 times more in prompt tokens alone, before the agent even starts working."
    )

    doc.add_heading("Request Concise Output", level=3)

    add_bullet(doc, "When you only need code, say so: \"return only the code, no explanation.\" This reduces output tokens significantly.")
    add_bullet(doc, "Use plan mode (available in both Claude Code and Codex) to have the agent think through the approach before writing code. This reduces iteration cycles and wasted output.")
    add_bullet(doc, "Batch related questions into a single structured prompt instead of asking them one at a time. Each round-trip consumes tokens for the repeated context.")

    doc.add_heading("Plan Your Workload", level=3)

    add_bullet(doc, "Run /check-usage (or check the extension dashboard) before starting heavy engineering days. If weekly usage is already above 75%, plan lighter tasks or use a cheaper model.", "Check before you start: ")
    add_bullet(doc, "Each 5-hour session window resets independently. Spreading a large project across multiple windows (morning and afternoon, or across days) gives you fresh quota for each block.", "Spread across windows: ")
    add_bullet(doc, "Schedule complex, token-heavy work (architecture design, comprehensive code review, multi-file refactors) for times when your limits are fresh. Save simple tasks for when limits are elevated.", "Schedule heavy work strategically: ")
    add_bullet(doc, "Configure the Claude Usage Monitor's auto-switcher to downgrade to a cheaper model before you hit hard limits. This prevents interruptions: you keep working (at slightly reduced capability) instead of being blocked entirely.", "Use the auto-switcher: ")

    # --- Urgency Reference ---
    doc.add_heading("Urgency Reference", level=2)

    doc.add_paragraph(
        "All three monitoring tools (the extension, the CLI hook, and the /check-usage command) use the same urgency classification. "
        "Use this table as a quick reference for what action to take at each level:"
    )

    add_styled_table(doc,
        ["Usage %", "Urgency", "Status", "Recommended Action"],
        [
            ["0 to 50%", "Low", "Green", "Continue freely. No action needed."],
            ["51 to 75%", "Moderate", "Yellow", "Be mindful. Consider switching to a cheaper model for simple tasks."],
            ["76 to 90%", "High", "Orange", "Switch to a cheaper model. Batch remaining work. Avoid exploratory prompts."],
            ["91 to 100%", "Critical", "Red", "Wait for the window to reset, switch to a different platform, or use the cheapest model available."],
        ]
    )

    doc.save(OUTPUT)
    print(f"Saved to: {OUTPUT}")


if __name__ == "__main__":
    build_document()
